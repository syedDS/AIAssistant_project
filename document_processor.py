"""
Document Processing Module
Extract text, process documents, and manage the data_store
Includes duplicate detection for both file tracker AND ChromaDB
"""
import os
import json
import hashlib
from datetime import datetime
from io import BytesIO

import PyPDF2
from docx import Document as DocxDocument
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter

from config import (
    UPLOAD_FOLDER, SUPPORTED_EXTENSIONS, 
    CHUNK_SIZE, CHUNK_OVERLAP, runtime_config
)


# Text splitter for chunking documents
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=CHUNK_SIZE, 
    chunk_overlap=CHUNK_OVERLAP
)


def get_document_hash(text: str) -> str:
    """Generate a hash of document content for duplicate detection"""
    return hashlib.md5(text.encode('utf-8')).hexdigest()


def is_document_in_chroma(filename: str, vector_store, content_hash: str = None) -> bool:
    """
    Check if document already exists in ChromaDB.
    Checks by filename in metadata OR by content hash.
    """
    try:
        # Method 1: Check by source filename in metadata
        # Get the underlying collection
        collection = vector_store._collection
        
        # Query for documents with this source
        results = collection.get(
            where={"source": filename},
            limit=1
        )
        
        if results and results.get('ids') and len(results['ids']) > 0:
            print(f"   📋 Found existing chunks for '{filename}' in ChromaDB")
            return True
        
        # Method 2: Check by document ID pattern
        # Our IDs are formatted as "filename_chunkindex"
        results = collection.get(
            ids=[f"{filename}_0"],  # Check if first chunk exists
            include=[]
        )
        
        if results and results.get('ids') and len(results['ids']) > 0:
            print(f"   📋 Found existing document ID for '{filename}' in ChromaDB")
            return True
            
    except Exception as e:
        # If there's an error, assume not indexed (safer to re-index)
        print(f"   ⚠️ ChromaDB check error: {e}")
    
    return False


def get_indexed_sources_from_chroma(vector_store) -> set:
    """
    Get all unique source filenames from ChromaDB.
    Useful for syncing with file tracker.
    """
    sources = set()
    
    try:
        collection = vector_store._collection
        
        # Get all documents (with pagination for large collections)
        batch_size = 1000
        offset = 0
        
        while True:
            results = collection.get(
                limit=batch_size,
                offset=offset,
                include=["metadatas"]
            )
            
            if not results or not results.get('metadatas'):
                break
            
            for metadata in results['metadatas']:
                if metadata and 'source' in metadata:
                    sources.add(metadata['source'])
            
            if len(results['ids']) < batch_size:
                break
            
            offset += batch_size
    
    except Exception as e:
        print(f"⚠️ Error getting sources from ChromaDB: {e}")
    
    return sources


def delete_document_from_chroma(filename: str, vector_store) -> int:
    """
    Delete all chunks of a document from ChromaDB.
    Returns number of chunks deleted.
    """
    deleted_count = 0
    
    try:
        collection = vector_store._collection
        
        # Find all chunk IDs for this document
        results = collection.get(
            where={"source": filename},
            include=[]
        )
        
        if results and results.get('ids'):
            ids_to_delete = results['ids']
            collection.delete(ids=ids_to_delete)
            deleted_count = len(ids_to_delete)
            print(f"   🗑️ Deleted {deleted_count} chunks for '{filename}' from ChromaDB")
    
    except Exception as e:
        print(f"   ⚠️ Error deleting from ChromaDB: {e}")
    
    return deleted_count


def extract_text_from_file(filepath):
    """Extract text from a file based on its extension"""
    file_ext = filepath.rsplit('.', 1)[-1].lower() if '.' in filepath else ''
    text = ""
    
    try:
        if file_ext == 'pdf':
            with open(filepath, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                text = "".join([page.extract_text() or "" for page in pdf_reader.pages])
        
        elif file_ext == 'docx':
            doc = DocxDocument(filepath)
            text = "\n".join([para.text for para in doc.paragraphs])
        
        elif file_ext == 'txt':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        
        elif file_ext == 'csv':
            with open(filepath, 'r', encoding='utf-8') as f:
                text = f.read()
        
        elif file_ext == 'json':
            with open(filepath, 'r', encoding='utf-8') as f:
                data = json.load(f)
                text = json.dumps(data, indent=2)
    
    except Exception as e:
        print(f"Error extracting text from {filepath}: {e}")
    
    return text


def extract_text_from_bytes(file_content, filename):
    """Extract text from file bytes (for uploads)"""
    file_ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
    text = ""
    
    try:
        if file_ext == 'pdf':
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            text = "".join([page.extract_text() or "" for page in pdf_reader.pages])
        
        elif file_ext == 'docx':
            doc = DocxDocument(BytesIO(file_content))
            text = "\n".join([para.text for para in doc.paragraphs])
        
        elif file_ext in ['txt', 'csv']:
            text = file_content.decode('utf-8')
        
        elif file_ext == 'json':
            data = json.loads(file_content.decode('utf-8'))
            text = json.dumps(data, indent=2)
    
    except Exception as e:
        print(f"Error extracting text from {filename}: {e}")
    
    return text


def process_document(text, filename, source_path, neo4j_graph, vector_store, 
                     entity_extractor=None, force_reindex=False):
    """
    Process document - create vector embeddings and optionally extract entities.
    
    Args:
        text: Document text content
        filename: Name of the file
        source_path: Full path to file
        neo4j_graph: Neo4j graph instance (can be None if KG disabled)
        vector_store: ChromaDB vector store
        entity_extractor: Entity extractor instance (can be None if extraction disabled)
        force_reindex: If True, delete existing and re-index
    """
    print(f"📄 Processing {filename}...")
    
    # Check if already in ChromaDB
    if not force_reindex and is_document_in_chroma(filename, vector_store):
        print(f"   ⏭️ Skipping - already indexed in ChromaDB")
        return {
            'entities_extracted': 0,
            'relationships_extracted': 0,
            'chunks_created': 0,
            'validation_errors': [],
            'skipped': True,
            'reason': 'already_in_chroma'
        }
    
    # If force reindex, delete existing first
    if force_reindex:
        delete_document_from_chroma(filename, vector_store)
    
    entity_count = 0
    rel_count = 0
    validation_errors = []
    
    # =========================================================================
    # OPTIONAL: Entity extraction (only if enabled and extractor provided)
    # =========================================================================
    if runtime_config.entity_extraction_enabled and entity_extractor and neo4j_graph:
        try:
            print(f"   🔍 Extracting entities (this may take a moment)...")
            extraction = entity_extractor.extract_with_validation(text, min_confidence=0.7)
            
            print(f"   📊 Extracted: {len(extraction['entities'])} entities, {len(extraction['relationships'])} relationships")
            
            # Add entities to Neo4j
            for entity in extraction.get('entities', []):
                canonical_id = neo4j_graph.add_validated_entity(
                    entity['name'],
                    entity['type'],
                    entity.get('properties', {}),
                    entity['confidence']
                )
                if canonical_id:
                    entity_count += 1
            
            # Add relationships to Neo4j
            for rel in extraction.get('relationships', []):
                source_entity = next(
                    (e for e in extraction['entities'] if e['name'] == rel['source']), 
                    None
                )
                target_entity = next(
                    (e for e in extraction['entities'] if e['name'] == rel['target']), 
                    None
                )
                
                if source_entity and target_entity:
                    success = neo4j_graph.add_validated_relationship(
                        rel['source'],
                        rel['target'],
                        source_entity['type'],
                        target_entity['type'],
                        rel['type'],
                        rel.get('properties', {}),
                        rel['confidence'],
                        rel.get('evidence', '')
                    )
                    if success:
                        rel_count += 1
            
            validation_errors = extraction.get('validation_errors', [])
            
        except Exception as e:
            print(f"   ⚠️ Entity extraction failed: {e}")
            validation_errors.append(str(e))
    else:
        print(f"   ⚡ Fast mode - skipping entity extraction")
    
    # =========================================================================
    # ALWAYS: Create vector chunks for ChromaDB
    # =========================================================================
    chunks = text_splitter.split_text(text)
    documents = []
    content_hash = get_document_hash(text)
    
    # Max characters per chunk (embedding model limit ~512 tokens ≈ 2000 chars)
    MAX_CHUNK_CHARS = 1500
    
    for i, chunk in enumerate(chunks):
        # Truncate if too long for embedding model
        if len(chunk) > MAX_CHUNK_CHARS:
            chunk = chunk[:MAX_CHUNK_CHARS] + "..."
            print(f"   ⚠️ Truncated chunk {i} (was too long)")
        
        doc = Document(
            page_content=chunk,
            metadata={
                'source': filename,
                'source_path': source_path or filename,
                'chunk_index': i,
                'total_chunks': len(chunks),
                'content_hash': content_hash,
                'indexed_at': datetime.now().isoformat()
            }
        )
        documents.append(doc)
    
    # Add to ChromaDB with unique IDs - batch to handle errors
    if documents:
        ids = [f"{filename}_{i}" for i in range(len(documents))]
        
        # Try batch first, fall back to individual if fails
        try:
            vector_store.add_documents(documents, ids=ids)
            print(f"   ✅ Added {len(documents)} chunks to ChromaDB")
        except Exception as batch_error:
            print(f"   ⚠️ Batch add failed: {batch_error}")
            print(f"   🔄 Trying individual chunks...")
            
            success_count = 0
            for doc, doc_id in zip(documents, ids):
                try:
                    # Further truncate if still failing
                    if len(doc.page_content) > 1000:
                        doc.page_content = doc.page_content[:1000] + "..."
                    vector_store.add_documents([doc], ids=[doc_id])
                    success_count += 1
                except Exception as e:
                    print(f"      ❌ Failed chunk {doc_id}: {str(e)[:50]}")
            
            print(f"   ✅ Added {success_count}/{len(documents)} chunks")
    
    return {
        'entities_extracted': entity_count,
        'relationships_extracted': rel_count,
        'chunks_created': len(documents),
        'validation_errors': validation_errors,
        'skipped': False,
        'content_hash': content_hash
    }


def sync_tracker_with_chroma(indexed_tracker, vector_store):
    """
    Sync file tracker with ChromaDB to ensure consistency.
    Adds any documents found in ChromaDB but missing from tracker.
    """
    print("🔄 Syncing file tracker with ChromaDB...")
    
    chroma_sources = get_indexed_sources_from_chroma(vector_store)
    tracker_files = set(os.path.basename(f) for f in indexed_tracker.get_all_indexed())
    
    # Find sources in ChromaDB but not in tracker
    missing_from_tracker = chroma_sources - tracker_files
    
    if missing_from_tracker:
        print(f"   Found {len(missing_from_tracker)} files in ChromaDB not in tracker")
        for filename in missing_from_tracker:
            # Add to tracker with minimal info
            indexed_tracker.indexed_files[filename] = {
                'mtime': 0,  # Unknown
                'indexed_at': 'synced_from_chroma',
                'stats': {'synced': True}
            }
        indexed_tracker.save()
    
    return len(missing_from_tracker)


def scan_and_index_data_store(neo4j_graph, vector_store, entity_extractor, indexed_tracker):
    """
    Scan data_store folder and index any unprocessed files.
    Called at startup to ensure all documents are searchable.
    
    Checks BOTH:
    1. File tracker (indexed_files.json)
    2. ChromaDB collection
    
    Only indexes if file is new to BOTH systems.
    """
    print(f"\n📂 Scanning data_store folder: {UPLOAD_FOLDER}")
    
    # First, sync tracker with ChromaDB
    sync_tracker_with_chroma(indexed_tracker, vector_store)
    
    # Get sources already in ChromaDB
    chroma_sources = get_indexed_sources_from_chroma(vector_store)
    print(f"   📊 Found {len(chroma_sources)} documents already in ChromaDB")
    
    files_indexed = 0
    files_skipped_tracker = 0
    files_skipped_chroma = 0
    
    for root, dirs, files in os.walk(UPLOAD_FOLDER):
        for filename in files:
            filepath = os.path.join(root, filename)
            file_ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''
            
            # Skip unsupported files
            if file_ext not in SUPPORTED_EXTENSIONS:
                continue
            
            # Check 1: Skip if in file tracker (and not modified)
            if indexed_tracker.is_indexed(filepath):
                files_skipped_tracker += 1
                continue
            
            # Check 2: Skip if already in ChromaDB
            if filename in chroma_sources:
                print(f"   ⏭️ {filename} - already in ChromaDB, updating tracker")
                indexed_tracker.mark_indexed(filepath, {'synced_from_chroma': True})
                files_skipped_chroma += 1
                continue
            
            # New file - index it
            print(f"   📄 Indexing NEW file: {filename}")
            
            try:
                text = extract_text_from_file(filepath)
                if text.strip():
                    result = process_document(
                        text, filename, filepath,
                        neo4j_graph, vector_store, entity_extractor
                    )
                    
                    if not result.get('skipped'):
                        indexed_tracker.mark_indexed(filepath, result)
                        files_indexed += 1
                        print(f"      ✅ {result['chunks_created']} chunks, {result['entities_extracted']} entities")
                    else:
                        files_skipped_chroma += 1
                else:
                    print(f"      ⚠️ No text extracted")
            except Exception as e:
                print(f"      ❌ Error: {e}")
    
    print(f"\n📊 data_store scan complete:")
    print(f"   - NEW files indexed: {files_indexed}")
    print(f"   - Skipped (in tracker): {files_skipped_tracker}")
    print(f"   - Skipped (in ChromaDB): {files_skipped_chroma}")
    print(f"   - Total tracked: {len(indexed_tracker.get_all_indexed())}")
    print(f"   - Total in ChromaDB: {len(chroma_sources) + files_indexed}")
    
    return files_indexed
